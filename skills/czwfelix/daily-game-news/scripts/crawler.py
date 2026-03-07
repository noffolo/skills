#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Daily Game News Crawler
每日游戏资讯自动抓取与报告生成
根据配置文件动态抓取所有配置的网站
"""

import json
import os
import re
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置路径
CONFIG_PATH = Path("/home/admin/.openclaw/workspace/configs/news-crawler-config.json")
REPORTS_DIR = Path("/home/admin/.openclaw/workspace/reports/daily-game-news")

# 确保报告目录存在
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def load_config():
    """加载配置文件"""
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_date_range():
    """获取 24 小时日期范围"""
    today = datetime.now()
    yesterday = today - timedelta(days=1)
    return yesterday.strftime('%Y%m%d'), today.strftime('%Y%m%d')


def fetch_with_searxng(site, date_start, date_end, limit=10):
    """使用 SearXNG 搜索抓取文章"""
    import subprocess
    
    # 构建搜索查询（不使用日期范围，SearXNG 不支持）
    search_query = f'site:{site}'
    
    try:
        # 使用相对路径
        script_path = os.path.join(os.path.dirname(__file__), '../../searxng/scripts/searxng.py')
        script_path = os.path.abspath(script_path)
        
        result = subprocess.run(
            ['uv', 'run', script_path,
             'search', search_query, '-n', str(limit), '--format', 'json'],
            capture_output=True, text=True, timeout=60,
            env={**os.environ, 'SEARXNG_URL': 'http://localhost:8080'},
            cwd=os.path.dirname(script_path)
        )
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            results = data.get('results', [])
            
            # 手动过滤 24 小时内的文章
            filtered = []
            today = datetime.now().date()
            yesterday = today - timedelta(days=1)
            
            for item in results:
                pub_date = item.get('publishedDate', '')
                if pub_date:
                    try:
                        # 解析日期
                        if 'T' in pub_date:
                            article_date = datetime.fromisoformat(pub_date.replace('Z', '+00:00')).date()
                        else:
                            article_date = datetime.strptime(pub_date[:10], '%Y-%m-%d').date()
                        
                        # 只保留今天和昨天的文章
                        if article_date >= yesterday:
                            filtered.append(item)
                    except:
                        # 日期解析失败，保留文章
                        filtered.append(item)
                else:
                    # 没有日期，保留
                    filtered.append(item)
            
            return filtered[:limit]
        else:
            print(f"  ⚠️ SearXNG 错误：{result.stderr[:200]}")
    except subprocess.TimeoutExpired:
        print(f"  ⚠️ SearXNG 超时")
    except Exception as e:
        print(f"  ⚠️ SearXNG 搜索失败：{e}")
    return []


def fetch_website(config, date_start, date_end):
    """根据配置抓取单个网站"""
    site_id = config['id']
    site_name = config['name']
    base_url = config['base_url']
    
    articles = []
    
    for section in config.get('分区', []):
        limit = section.get('筛选数量', 2)
        section_name = section.get('name', '')
        
        print(f"  - 抓取 {site_name} - {section_name}..." )
        
        # 根据网站选择不同的抓取策略
        if site_id == 'gcores':
            results = fetch_with_searxng('gcores.com/articles', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', '').replace(' - 机核 GCORES', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'gamersky':
            results = fetch_with_searxng('gamersky.com/news', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', '').replace(' _ 游民星空 GamerSky.com', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'yystv':
            results = fetch_with_searxng('yystv.cn/p', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', '').replace(' - 游研社', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'ign':
            results = fetch_with_searxng('ign.com/articles', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', '').replace(' - IGN', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'youxituoluo':
            results = fetch_with_searxng('youxituoluo.com', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'chuapp':
            results = fetch_with_searxng('chuapp.com', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', '').replace(' - 触乐', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'gamelook':
            results = fetch_with_searxng('gamelook.com.cn', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'gamespot':
            results = fetch_with_searxng('gamespot.com/news', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', '').replace(' - GameSpot', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        elif site_id == 'gamedeveloper':
            results = fetch_with_searxng('gamedeveloper.com', date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        else:
            # 通用抓取方法
            domain = base_url.replace('https://', '').replace('http://', '').split('/')[0]
            results = fetch_with_searxng(domain, date_start, date_end, limit * 2)
            for item in results[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', '') or '今日',
                    'content': item.get('content', '')[:200]
                })
        
        print(f"    ✓ 抓取 {len(articles)} 篇")
    
    return articles


def classify_article(title, content, config):
    """根据配置自动分类文章"""
    text = (title + ' ' + content).lower()
    
    # 从配置中读取分类规则
    rules = config.get('分类规则', {}).get('规则', [])
    
    for rule in rules:
        condition = rule.get('条件', '')
        category = rule.get('分类', '其他')
        
        # 简单解析条件（支持 OR 和 AND）
        if ' OR ' in condition:
            parts = condition.split(' OR ')
            for part in parts:
                part = part.strip().strip("'\"")
                if '|' in part:
                    # 正则匹配
                    keywords = part.split('|')
                    for kw in keywords:
                        if kw.lower().strip() in text:
                            return category
                elif part.lower() in text:
                    return category
        elif ' AND ' in condition:
            parts = condition.split(' AND ')
            match_all = True
            for part in parts:
                part = part.strip().strip("'\"")
                if '|' in part:
                    keywords = part.split('|')
                    if not any(kw.lower().strip() in text for kw in keywords):
                        match_all = False
                        break
                elif part.lower() not in text:
                    match_all = False
                    break
            if match_all:
                return category
        else:
            # 简单条件
            if '|' in condition:
                keywords = condition.split('|')
                for kw in keywords:
                    if kw.lower().strip().strip("'\"") in text:
                        return category
            elif condition.lower().strip("'\"") in text:
                return category
    
    return '其他'


def generate_markdown_report(articles_by_category, date_str, config):
    """生成 Markdown 格式报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report = f"""# 📰 每日游戏资讯报告

**报告日期**: {date_str}  
**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**数据来源**: {', '.join([site['name'] for site in config.get('网站配置', [])])}

---

"""
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            report += f"## {emoji} {category}\n\n"
            report += "| 发布源 | 时间 | 文章标题 | 链接 |\n"
            report += "|--------|------|----------|------|\n"
            
            for article in articles:
                title = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                published = article['published'][:16] if article['published'] else '今日'
                report += f"| {article['source']} | {published} | {title} | [查看详情]({article['url']}) |\n"
            
            report += "\n---\n\n"
    
    report += f"\n**报告结束** - 共抓取 {total_articles} 篇文章\n"
    
    return report


def generate_word_report(articles_by_category, date_str, config):
    """生成 Word 格式报告 (.docx)"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    # 创建 Word 文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('📰 每日游戏资讯报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph(f'报告日期：{date_str}')
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据来源：{", ".join([site["name"] for site in config.get("网站配置", [])])}')
    doc.add_paragraph()
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            # 分类标题
            heading = doc.add_heading(f'{emoji} {category}', level=1)
            heading.runs[0].bold = True
            
            # 创建表格
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # 表头
            header_cells = table.rows[0].cells
            headers = ['发布源', '时间', '文章标题', '链接']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # 添加文章
            for article in articles:
                row = table.add_row().cells
                row[0].text = article['source']
                published = article['published'][:16] if article['published'] else '今日'
                row[1].text = published
                title_text = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                row[2].text = title_text
                row[3].text = article['url']
            
            doc.add_paragraph()
    
    # 总结
    doc.add_paragraph()
    doc.add_paragraph(f'本报告共抓取 {total_articles} 篇文章').bold = True
    
    # 保存文件
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.docx'
    doc.save(file_path)
    
    return file_path


def generate_text_report(articles_by_category, date_str, config):
    """生成文本格式报告（保留用于兼容）"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report_lines = []
    report_lines.append("=" * 60)
    report_lines.append("📰 每日游戏资讯报告")
    report_lines.append("=" * 60)
    report_lines.append(f"报告日期：{date_str}")
    report_lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append(f"数据来源：{', '.join([site['name'] for site in config.get('网站配置', [])])}")
    report_lines.append("")
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            report_lines.append(f"\n{emoji} {category}")
            report_lines.append("-" * 40)
            
            for i, article in enumerate(articles, 1):
                report_lines.append(f"\n{i}. {article['title']}")
                report_lines.append(f"   发布源：{article['source']}")
                published = article['published'][:16] if article['published'] else '今日'
                report_lines.append(f"   时间：{published}")
                report_lines.append(f"   链接：{article['url']}")
                if article.get('content'):
                    report_lines.append(f"   摘要：{article['content']}...")
    
    report_lines.append("\n" + "=" * 60)
    report_lines.append(f"本报告共抓取 {total_articles} 篇文章")
    report_lines.append("=" * 60)
    
    # 保存为文本文件
    report_text = "\n".join(report_lines)
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.txt'
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(report_text)
    
    return file_path


def main():
    """主函数"""
    print("=" * 60)
    print("📰 Daily Game News Crawler")
    print("每日游戏资讯自动抓取")
    print("=" * 60)
    
    # 加载配置
    print("\n📋 加载配置文件...")
    config = load_config()
    print(f"  ✓ 配置文件：{CONFIG_PATH}")
    print(f"  ✓ 配置网站数量：{len(config.get('网站配置', []))}")
    
    # 获取日期范围
    date_start, date_end = get_date_range()
    date_str = datetime.now().strftime('%Y-%m-%d')
    print(f"\n📅 抓取日期范围：{date_start} - {date_end}")
    print(f"  报告日期：{date_str}")
    
    # 抓取所有网站
    print("\n🕷️ 开始抓取文章...")
    all_articles = []
    
    for site_config in config.get('网站配置', []):
        site_name = site_config.get('name', 'Unknown')
        priority = site_config.get('抓取优先级', 2)
        
        # 只抓取优先级为 1 的网站
        if priority == 1:
            print(f"\n【{site_name}】")
            articles = fetch_website(site_config, date_start, date_end)
            
            # 分类文章
            for article in articles:
                category = classify_article(article['title'], article.get('content', ''), config)
                article['category'] = category
            
            all_articles.extend(articles)
        else:
            print(f"\n⏭️ 跳过 {site_name} (优先级：{priority})")
    
    print(f"\n✅ 共抓取 {len(all_articles)} 篇文章")
    
    # 按分类整理
    print("\n🏷️ 分类整理...")
    articles_by_category = {}
    for article in all_articles:
        category = article.get('category', '其他')
        if category not in articles_by_category:
            articles_by_category[category] = []
        articles_by_category[category].append(article)
        print(f"  - {article['title'][:30]}... → {category}")
    
    # 生成报告
    print("\n📝 生成报告...")
    
    # Markdown 报告
    markdown_report = generate_markdown_report(articles_by_category, date_str, config)
    print("  ✓ Markdown 报告生成完成")
    
    # Word 报告
    word_path = generate_word_report(articles_by_category, date_str, config)
    print(f"  ✓ Word 报告已保存：{word_path}")
    
    # 文本报告（保留用于兼容）
    text_path = generate_text_report(articles_by_category, date_str, config)
    print(f"  ✓ 文本报告已保存：{text_path}")
    
    # 打印报告预览
    print("\n" + "=" * 60)
    print("📊 报告预览")
    print("=" * 60)
    print(markdown_report[:2000])
    
    print("\n" + "=" * 60)
    print("✅ 报告生成完成！")
    print("=" * 60)
    print(f"\n📁 Word 报告位置：{word_path}")
    print(f"📁 文本报告位置：{text_path}")
    print(f"\n💡 获取报告指令:")
    print(f"   read {word_path}")
    
    return {
        'date': date_str,
        'total_articles': len(all_articles),
        'categories': articles_by_category,
        'word_path': str(word_path)
    }


if __name__ == '__main__':
    main()
